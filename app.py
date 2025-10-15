# -*- coding: utf-8 -*-
import os
import sys
import io
import re
import json
import base64
import shutil
import tempfile
import subprocess
import time
import threading
import logging
from logging.handlers import RotatingFileHandler
from pathlib import Path
from datetime import datetime, timedelta
import zipfile
from docx2pdf import convert as docx2pdf_convert
import mysql.connector
import html
import schedule
from flask import Flask, request, jsonify, send_file, g
from flask_cors import CORS
from bs4 import BeautifulSoup, NavigableString
import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import fitz  # PyMuPDF

from dotenv import load_dotenv
load_dotenv()

# === LOGGING ===
def _setup_logging():
    """Configura logging com dois arquivos: INFO e WARNING+/ERROR."""
    log_dir = os.getenv("LOG_DIR")
    os.makedirs(log_dir, exist_ok=True)

    logger = logging.getLogger("ridolf_app")
    logger.setLevel(logging.INFO)
    logger.propagate = False  # não repassar para o root

    fmt = logging.Formatter(
        fmt="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S"
    )

    class MaxLevelFilter(logging.Filter):
        """Permite no handler apenas mensagens até um dado nível (inclusive)."""
        def __init__(self, max_level):
            super(MaxLevelFilter, self).__init__()
            self.max_level = max_level
        def filter(self, record):
            return record.levelno <= self.max_level

    # INFO handler (somente INFO, sem WARNING+)
    info_path = os.path.join(log_dir, "app_info.log")
    h_info = RotatingFileHandler(info_path, maxBytes=5*1024*1024, backupCount=5, encoding="utf-8")
    h_info.setLevel(logging.INFO)
    h_info.addFilter(MaxLevelFilter(logging.INFO))
    h_info.setFormatter(fmt)

    # ERROR/ALERT handler (WARNING e acima)
    err_path = os.path.join(log_dir, "app_errors.log")
    h_err = RotatingFileHandler(err_path, maxBytes=5*1024*1024, backupCount=5, encoding="utf-8")
    h_err.setLevel(logging.WARNING)
    h_err.setFormatter(fmt)

    logger.handlers = [h_info, h_err]
    return logger

logger = _setup_logging()

# === CONFIG / DB / HELPERS ===
db_config = dict(
    host = os.getenv("DB_HOST"),
    port = os.getenv("DB_PORT"),
    user = os.getenv("DB_USER"),
    password = os.getenv("DB_PASSWORD"),
    database = os.getenv("DB_NAME"),
    auth_plugin = os.getenv("DB_AUTH_PLUGIN"),
    charset = "utf8mb4",
    use_unicode = True,
)

def get_connection():
    """
    Abre e retorna uma conexão MySQL baseada em db_config, com validações,
    timeout e logs — sem derrubar o processo em caso de erro.
    """
    cfg = db_config.copy()

    # Remova chaves None/vazias (evita bugs do connector)
    cfg = {k: v for k, v in cfg.items() if v not in (None, "", "None")}

    # Campos obrigatórios
    obrig = ("host", "user", "password", "database")
    faltando = [k for k in obrig if k not in cfg]
    if faltando:
        msg = f"Variáveis de ambiente do DB faltando: {faltando}"
        error_logger.error(msg)
        raise RuntimeError(msg)

    # Porta como int (se vier string)
    if "port" in cfg:
        try:
            cfg["port"] = int(cfg["port"])
        except Exception:
            error_logger.warning("DB_PORT inválida (%r); removendo.", cfg["port"])
            cfg.pop("port", None)

    # Evita crash do connector com auth_plugin inválido
    if "auth_plugin" in cfg and cfg["auth_plugin"].lower() not in (
        "mysql_native_password", "caching_sha2_password"
    ):
        error_logger.warning("auth_plugin desconhecido (%r); removendo.", cfg["auth_plugin"])
        cfg.pop("auth_plugin", None)

    # Força timeout e modo puro python
    cfg.setdefault("connection_timeout", 10)
    cfg.setdefault("autocommit", True)
    cfg.setdefault("use_pure", True)  # evita extensões C

    try:
        cnx = mysql.connector.connect(**cfg)
        return cnx
    except mysql.connector.Error as e:
        error_logger.error("MySQL connector error: %s", e, exc_info=True)
        raise
    except Exception as e:
        error_logger.error("Falha inesperada ao conectar MySQL: %s", e, exc_info=True)
        raise

def executar(sql, params=()):
    """Executa INSERT/UPDATE/DELETE simples e retorna lastrowid (quando houver)."""
    with get_connection() as cnx, cnx.cursor() as cur:
        cur.execute(sql, params)
        cnx.commit()
        return cur.lastrowid

def nome_unico(cnx, caminho, base_original):
    """Gera nome único dentro de um caminho, adicionando sufixo (n) quando necessário."""
    base = base_original
    query = """
        SELECT nome_arquivo
          FROM galeria_juridico
         WHERE caminho = %s
           AND NOT (
               COALESCE(status,'') = 'deletado'
               OR (status = 'alterado' AND alteracao = 'deletado')
           )
           AND nome_arquivo LIKE %s
    """
    with cnx.cursor() as cur:
        cur.execute(query, (caminho, f"{base}%"))
        existentes = {row[0] for row in cur.fetchall()}

    if base not in existentes:
        return base

    padrao = re.compile(rf"^{re.escape(base)}\((\d+)\)$")
    max_i  = max(
        (int(m.group(1)) for n in existentes if (m := padrao.match(n))),
        default = 0
    )
    return f"{base}({max_i + 1})"

def get_user_by_id(id_usuario):
    """Recupera o registro de usuário (tabela_funcionarios) pelo ID."""
    cnx = get_connection()
    cur = cnx.cursor(dictionary=True)
    cur.execute("SELECT * FROM tabela_funcionarios WHERE idtabela_funcionarios = %s",(id_usuario,),)
    row = cur.fetchone()
    cur.close()
    cnx.close()
    return row

def run_pandoc_for_plain_text(bytes_docx):
    """Extrai texto 'plain' do DOCX com pandoc (ajuda a reconstruir marcadores numéricos)."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_in:
            temp_in.write(bytes_docx)
            temp_in_path = temp_in.name

        comando = ["pandoc", temp_in_path, "--from", "docx", "--to", "plain", "--wrap=none"]
        resultado = subprocess.run(comando, capture_output=True, text=True, check=True, encoding='utf-8')
        return [line.strip() for line in resultado.stdout.split('\n') if line.strip()]
    except Exception:
        logger.exception("run_pandoc_for_plain_text: falha na execução do pandoc.")
        return []
    finally:
        try:
            if 'temp_in_path' in locals() and os.path.exists(temp_in_path):
                os.remove(temp_in_path)
        except Exception:
            logger.warning("run_pandoc_for_plain_text: falha ao remover temporário.")

def merge_split_paragraphs(html_content):
    """Une apenas parágrafos quebrados artificialmente, preservando blocos válidos e tabelas/imagens."""
    if not html_content or not html_content.strip():
        return ""

    soup = BeautifulSoup(f'<body>{html_content}</body>', 'html.parser')
    body = soup.body
    blocks = list(body.find_all(['p', 'table', 'img'], recursive=False))
    final_elems = []

    def _no_margin(p):
        return 'margin-bottom' not in (p.get('style') or '')

    def _text(p):
        return (p.get_text(" ", strip=True) or '')

    def _starts_list_or_clause(s):
        return bool(re.match(r'^(Art\.\s*\d+|§\s*\d+[ºª]?|[IVXLCDM]+\.\s|[A-Za-z][\.\)]\s|\d+\.\s)', s))

    def _looks_heading(s):
        s_norm = re.sub(r'\s+', ' ', s).strip()
        return len(s_norm) <= 120 and s_norm.isupper() and re.search(r'[A-Z]{3}', s_norm) is not None

    def _ends_sentence(s):
        return bool(re.search(r'[.!?;:»”\'")\]]\s*$', s))

    def _should_merge(a, b):
        if a.name != 'p' or b.name != 'p':
            return False
        if not _no_margin(a) or not _no_margin(b):
            return False

        ta, tb = _text(a), _text(b)
        if not ta or not tb:
            return False
        if _ends_sentence(ta):
            return False
        if _starts_list_or_clause(tb) or _looks_heading(tb):
            return False
        if len(ta) > 300 or len(tb) > 300:
            return False
        return True

    i = 0
    while i < len(blocks):
        cur = blocks[i]
        if cur.name != 'p':
            final_elems.append(cur.extract())
            i += 1
            continue

        if i + 1 < len(blocks) and _should_merge(cur, blocks[i + 1]):
            target = cur.extract()
            j = i + 1
            while j < len(blocks) and _should_merge(target, blocks[j]):
                nxt = blocks[j].extract()
                target.append(NavigableString(' '))
                for child in list(nxt.children):
                    target.append(child.extract())
                j += 1
            final_elems.append(target)
            i = j
        else:
            final_elems.append(cur.extract())
            i += 1

    return "".join(str(el) for el in final_elems)

def iter_block_items(parent, Document_class):
    """Itera blocos de um documento/célula (parágrafos e tabelas)."""
    if isinstance(parent, Document_class):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Parent type not supported")

    for child in parent_elm:
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield Table(child, parent)

def find_nearest(value, choices):
    """Retorna o valor em 'choices' mais próximo de 'value'."""
    return min(choices, key=lambda x: abs(x - value))

def process_paragraph(para, image_map, plain_text_ref):
    """Converte um parágrafo do DOCX para HTML preservando ênfases, imagens e quebras."""
    allowed_spacing_pts = [0, 5, 8, 12, 15, 18, 22, 25, 28, 32, 35, 38, 42, 45, 48, 52]
    styles = []

    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        styles.append('text-align: center;')
    elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        styles.append('text-align: right;')
    elif para.alignment in (
        WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.DISTRIBUTE,
        WD_ALIGN_PARAGRAPH.JUSTIFY_LOW, WD_ALIGN_PARAGRAPH.JUSTIFY_MED,
        WD_ALIGN_PARAGRAPH.JUSTIFY_HI, WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    ):
        styles.append('text-align: justify;')

    if para.paragraph_format and para.paragraph_format.space_after and hasattr(para.paragraph_format.space_after, 'pt'):
        pt_value = para.paragraph_format.space_after.pt
        if pt_value is not None and pt_value > 0:
            final_pt = find_nearest(pt_value, allowed_spacing_pts)
            if final_pt > 0:
                styles.append(f'margin-bottom: {final_pt}pt;')

    p_style_attr = f' style="{" ".join(styles)}"' if styles else ''
    p_html = ''
    marker_html = ''

    if plain_text_ref:
        regex = r'^(Art\.\s*[\d]+\.|§\s*\d+[ºª]|\b[IVXLCDM]+\.|[a-zA-Z][\.\)]|\d+\.)\s*'
        match = re.match(regex, plain_text_ref, re.IGNORECASE)
        if match:
            marker = match.group(1).strip()
            docx_text_clean = para.text.strip()
            if not docx_text_clean.lower().startswith(marker.lower()):
                marker_html = f'{marker}&emsp;'

    for run in para.runs:
        if '<a:blip' in run.element.xml:
            for r_id in image_map:
                if r_id in run.element.xml:
                    image_part = image_map[r_id]
                    image_bytes = image_part.blob
                    b64_string = base64.b64encode(image_bytes).decode('utf-8')
                    width_px, height_px = None, None
                    try:
                        extent = run.element.xpath('.//wp:extent')
                        if extent:
                            cx = int(extent[0].get('cx'))
                            cy = int(extent[0].get('cy'))
                            width_px = int(cx / 9525)
                            height_px = int(cy / 9525)
                    except Exception:
                        pass
                    style = f'width:{width_px}px; height:{height_px}px;' if width_px and height_px else 'max-width:100%;'
                    p_html += f'<img src="data:{image_part.content_type};base64,{b64_string}" style="{style}">'
                    break
            try:
                ns = run.element.nsmap
                br_qtd = len(run.element.xpath('.//w:br', namespaces=ns))
                tab_qtd = len(run.element.xpath('.//w:tab', namespaces=ns))
                if tab_qtd:
                    p_html += '&emsp;' * tab_qtd
                if br_qtd:
                    p_html += '<br/>' * br_qtd
            except Exception:
                pass
            continue

        text = run.text or ""
        if not text and '<w:br' not in run.element.xml and '<w:tab' not in run.element.xml:
            continue

        escaped_text = html.escape(text).replace('\t', '&emsp;')
        run_styles = []

        if run.font.size and hasattr(run.font.size, 'pt') and run.font.size.pt is not None:
            pt_value = run.font.size.pt
            rounded_pt = round(pt_value)
            final_pt = max(8, min(20, rounded_pt))
            run_styles.append(f'font-size: {final_pt}pt;')

        if run.font.name:
            run_styles.append(f"font-family: '{run.font.name}';")

        inner_html = escaped_text

        effective_bold = run.bold if run.bold is not None else para.style.font.bold
        effective_italic = run.italic if run.italic is not None else para.style.font.italic
        effective_underline = run.underline if run.underline is not None else para.style.font.underline

        if effective_underline:
            inner_html = f'<u>{inner_html}</u>'
        if effective_italic:
            inner_html = f'<em>{inner_html}</em>'
        if effective_bold:
            inner_html = f'<strong>{inner_html}</strong>'

        if run_styles:
            style_attr = "".join(run_styles)
            p_html += f'<span style="{style_attr}">{inner_html}</span>'
        else:
            p_html += inner_html

        try:
            ns = run.element.nsmap
            tab_qtd = len(run.element.xpath('.//w:tab', namespaces=ns))
            br_qtd  = len(run.element.xpath('.//w:br',  namespaces=ns))
            if tab_qtd:
                p_html += '&emsp;' * tab_qtd
            if br_qtd:
                p_html += '<br/>' * br_qtd
        except Exception:
            pass

    final_p_html = marker_html + p_html
    if final_p_html.strip():
        return f'<p{p_style_attr}>{final_p_html}</p>'
    return ''

def process_table(table, image_map, Document_class):
    """Converte uma tabela do DOCX para HTML, preservando parágrafos internos."""
    table_html = '<table border="1" style="width:100%; border-collapse: collapse;">'
    for row in table.rows:
        table_html += '<tr>'
        for cell in row.cells:
            cell_content = ''
            for block in iter_block_items(cell, Document_class):
                if isinstance(block, Paragraph):
                    cell_content += process_paragraph(block, image_map, '')
                elif isinstance(block, Table):
                    cell_content += process_table(block, image_map, Document_class)
            table_html += f'<td style="padding: 4px;">{cell_content}</td>'
        table_html += '</tr>'
    table_html += '</table>'
    return table_html

def docx_to_html(bytes_docx):
    """Converte DOCX em HTML concatenado (parágrafos, tabelas e imagens)."""
    doc = Document(io.BytesIO(bytes_docx))
    Document_class = type(doc)

    plain_text_paragraphs = run_pandoc_for_plain_text(bytes_docx)
    text_search_idx = 0

    image_map = {
        r.rId: r.target_part
        for r in doc.part.rels.values() if "image" in r.target_ref
    }

    html_body = ''
    for block in iter_block_items(doc, Document_class):
        if isinstance(block, Paragraph):
            docx_text = re.sub(r'\s+', ' ', block.text.strip())
            ref_text = ''
            if docx_text:
                for i in range(text_search_idx, len(plain_text_paragraphs)):
                    plain_text_line = re.sub(r'\s+', ' ', plain_text_paragraphs[i].strip())
                    if plain_text_line.endswith(docx_text):
                        ref_text = plain_text_paragraphs[i]
                        text_search_idx = i + 1
                        break
            html_body += process_paragraph(block, image_map, ref_text)
        elif isinstance(block, Table):
            html_body += process_table(block, image_map, Document_class)

    return f'{merge_split_paragraphs(html_body)}'

def parse_style(style_str):
    """Converte string de estilos inline CSS em dict {prop: valor}."""
    styles = {}
    if not style_str:
        return styles
    for rule in style_str.split(';'):
        if ':' in rule:
            key, value = rule.split(':', 1)
            styles[key.strip().lower()] = value.strip()
    return styles

def apply_text_formatting(text_node, run, stop_element):
    """Aplica ênfases e fonte em um run do DOCX com base nos elementos HTML ascendentes."""
    is_bold = False
    is_italic = False
    is_underline = False

    parent = text_node.find_parent()
    while parent is not None and parent != stop_element:
        if parent.name in ['strong', 'b']:
            is_bold = True
        if parent.name in ['em', 'i']:
            is_italic = True
        if parent.name == 'u':
            is_underline = True
        if parent.name == 'span':
            span_styles = parse_style(parent.get('style', ''))
            font_size_str = span_styles.get('font-size')
            if font_size_str and 'pt' in font_size_str:
                try:
                    size_val = float(re.sub(r'[^0-9.]', '', font_size_str))
                    run.font.size = Pt(size_val)
                except Exception:
                    pass
            font_family_str = span_styles.get('font-family')
            if font_family_str:
                font_name = font_family_str.strip().strip("'\"")
                run.font.name = font_name
                r = run._element
                r.rPr.rFonts.set(qn('w:ascii'), font_name)
                r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
                r.rPr.rFonts.set(qn('w:cs'), font_name)
        parent = parent.find_parent()

    run.bold = is_bold
    run.italic = is_italic
    run.underline = is_underline

def convert_html_to_docx(html_content):
    """Converte HTML (com estilos inline) em um arquivo DOCX (BytesIO)."""
    def _pt_from_css(val):
        if not val:
            return None
        s = str(val).strip().lower()
        m = re.search(r'(-?\d+(?:\.\d+)?)', s)
        if not m:
            return None
        x = float(m.group(1))
        if 'px' in s:
            return x * 72.0 / 96.0
        return x

    def _apply_paragraph_css(p_styles, p_format, p_docx=None):
        ta = (p_styles.get('text-align') or '').strip()
        if ta == 'center':
            p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif ta == 'right':
            p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif ta == 'justify':
            p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        mt = _pt_from_css(p_styles.get('margin-top'))
        mb = _pt_from_css(p_styles.get('margin-bottom'))
        ml = _pt_from_css(p_styles.get('margin-left'))
        mr = _pt_from_css(p_styles.get('margin-right'))
        ti = _pt_from_css(p_styles.get('text-indent'))

        if mt is not None: p_format.space_before = Pt(mt)
        if mb is not None: p_format.space_after  = Pt(mb)
        if mr is not None: p_format.right_indent = Pt(mr)

        if ti is not None:
            if ti < 0:
                hang = abs(ti)
                left = max(0.0, (ml or 0.0) - hang)
                if left > 0:
                    p_format.left_indent = Pt(left)
                p_format.first_line_indent = Pt(-hang)
            else:
                if ml is not None:
                    p_format.left_indent = Pt(ml)
                p_format.first_line_indent = Pt(ti)
        else:
            if ml is not None:
                p_format.left_indent = Pt(ml)

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    r = font._element
    r.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    r.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    r.rPr.rFonts.set(qn('w:cs'), 'Arial')

    pf = style.paragraph_format
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width  = Mm(210)
    section.left_margin = Mm(18.52)
    section.right_margin = Mm(18.52)
    section.top_margin = Mm(18.52)
    section.bottom_margin = Mm(18.52)

    soup = BeautifulSoup(html_content or "", "html.parser")
    body = soup.find('body') or soup

    def _emit_children_into_paragraph(parent_el, p_docx, stop_element):
        for child in parent_el.children:
            if getattr(child, 'name', None) == 'img':
                src = child.get('src', '')
                if 'data:image' in src and 'base64,' in src:
                    try:
                        img_data = src.split('base64,')[1]
                        img_bytes = base64.b64decode(img_data)
                        img_stream = io.BytesIO(img_bytes)
                        run = p_docx.add_run()
                        img_styles = parse_style(child.get('style', ''))
                        width_css = img_styles.get('width')
                        if width_css:
                            wpt = _pt_from_css(width_css) or 0
                            run.add_picture(img_stream, width=Inches(wpt / 72.0))
                        else:
                            run.add_picture(img_stream)
                    except Exception as e:
                        logger.warning("convert_html_to_docx: falha ao processar imagem aninhada: %s", e)

            elif getattr(child, 'name', None) == 'br':
                p_docx.add_run().add_break()

            elif isinstance(child, NavigableString):
                s = str(child)
                if s:
                    run = p_docx.add_run(s)
                    apply_text_formatting(child, run, stop_element)

            elif getattr(child, 'name', None):
                for text_node in child.find_all(string=True, recursive=True):
                    s = str(text_node)
                    if s:
                        run = p_docx.add_run(s)
                        apply_text_formatting(text_node, run, stop_element)

    def _emit_table(table_el):
        rows = table_el.find_all('tr')
        if not rows:
            return
        cols = len(rows[0].find_all(['td', 'th']))
        if cols == 0:
            return
        t = doc.add_table(rows=len(rows), cols=cols)
        t.style = 'Table Grid'
        for i, row_html in enumerate(rows):
            cells_html = row_html.find_all(['td', 'th'])
            for j, cell_html in enumerate(cells_html):
                cell_docx = t.cell(i, j)
                cell_docx.text = ''
                for p_in_cell in cell_html.find_all('p'):
                    p_docx_in_cell = cell_docx.add_paragraph()
                    pf2 = p_docx_in_cell.paragraph_format
                    pf2.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    pf2.space_before = Pt(0)
                    pf2.space_after = Pt(0)
                    pf2.first_line_indent = Pt(0)
                    _apply_paragraph_css(parse_style(p_in_cell.get('style','')), pf2, p_docx_in_cell)

                    txt = p_in_cell.get_text(strip=True)
                    if re.match(r'^\s*(\d+\.|[a-zA-Z][\.\)]|[IVXLCDM]+\.)', txt):
                        p_docx_in_cell.add_run('\u200b')

                    _emit_children_into_paragraph(p_in_cell, p_docx_in_cell, p_in_cell)

    for el in body.find_all(['p', 'table', 'img'], recursive=False):
        if el.name == 'p':
            p = doc.add_paragraph()
            pf3 = p.paragraph_format
            pf3.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf3.space_before = Pt(0)
            pf3.space_after = Pt(0)
            pf3.first_line_indent = Pt(0)
            pf3.left_indent = Pt(0)
            pf3.right_indent = Pt(0)

            _apply_paragraph_css(parse_style(el.get('style','')), pf3, p)

            txt = el.get_text(strip=True)
            if re.match(r'^\s*(\d+\.|[a-zA-Z][\.\)]|[IVXLCDM]+\.)', txt):
                p.add_run('\u200b')

            _emit_children_into_paragraph(el, p, el)

        elif el.name == 'table':
            _emit_table(el)

        elif el.name == 'img':
            src = el.get('src', '')
            if 'data:image' in src and 'base64,' in src:
                try:
                    img_data = src.split('base64,')[1]
                    img_bytes = base64.b64decode(img_data)
                    img_stream = io.BytesIO(img_bytes)
                    p_for_img = doc.add_paragraph()
                    p_for_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_for_img.add_run()
                    width_css = parse_style(el.get('style','')).get('width')
                    if width_css:
                        wpt = _pt_from_css(width_css) or 0
                        run.add_picture(img_stream, width=Inches(wpt / 72.0))
                    else:
                        run.add_picture(img_stream)
                except Exception as e:
                    logger.warning("convert_html_to_docx: falha ao processar imagem raiz: %s", e)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def _find_soffice():
    """Localiza o executável do LibreOffice/soffice para conversão DOCX→PDF."""
    env = os.environ.get("SOFFICE_PATH") or os.environ.get("LIBREOFFICE_PATH")
    candidates = []
    if env:
        candidates.append(env)

    if sys.platform.startswith("win"):
        candidates += [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
    else:
        exe = shutil.which("soffice") or shutil.which("libreoffice")
        if exe:
            candidates.append(exe)
        candidates += ["/usr/bin/soffice", "/usr/local/bin/soffice", "/snap/bin/libreoffice"]

    for c in candidates:
        if c and os.path.exists(c):
            return c
    return shutil.which("soffice") or shutil.which("libreoffice")

def _lo_user_install_arg(profile_dir):
    """Monta o argumento -env:UserInstallation=file:///... para evitar lock de perfil do LO."""
    uri = Path(profile_dir).absolute().as_posix()
    if not uri.startswith("file:///"):
        uri = f"file:///{uri}"
    return f"-env:UserInstallation={uri}"

def _docx_bytes_to_pdf_bytes(docx_bytes_io):
    """Converte BytesIO DOCX em BytesIO PDF via LibreOffice (fallback docx2pdf)."""
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "documento.docx")
        out_dir = td
        with open(in_path, "wb") as f:
            f.write(docx_bytes_io.getvalue())

        soffice = _find_soffice()
        last_err = None

        if soffice:
            def _try_lo_convert(filter_token, profile_name):
                profile_dir = os.path.join(td, profile_name)
                os.makedirs(profile_dir, exist_ok=True)

                creationflags = 0
                startupinfo = None
                if sys.platform.startswith("win"):
                    creationflags = 0x08000000  # CREATE_NO_WINDOW
                    startupinfo = subprocess.STARTUPINFO()
                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

                env = os.environ.copy()
                env["HOME"] = profile_dir
                env["USERPROFILE"] = profile_dir

                cmd = [
                    soffice,
                    "--headless",
                    "--invisible",
                    "--norestore",
                    "--nodefault",
                    "--nofirststartwizard",
                    "--nolockcheck",
                    _lo_user_install_arg(profile_dir),
                    "--convert-to", filter_token,
                    "--outdir", out_dir,
                    in_path
                ]
                res = subprocess.run(
                    cmd,
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    encoding="utf-8",
                    creationflags=creationflags,
                    startupinfo=startupinfo,
                    env=env
                )
                if not any(name.lower().endswith(".pdf") for name in os.listdir(out_dir)):
                    raise subprocess.CalledProcessError(
                        returncode=0, cmd=cmd, output=res.stdout, stderr=res.stderr
                    )

            try:
                _try_lo_convert("pdf:writer_pdf_Export", "lo_profile_1")
            except subprocess.CalledProcessError as e1:
                last_err = f"[LO#1] exit={e1.returncode} STDOUT={str(e1.stdout)[:1000]} STDERR={str(e1.stderr)[:1000]}"
                try:
                    _try_lo_convert("pdf", "lo_profile_2")
                    last_err = None
                except subprocess.CalledProcessError as e2:
                    last_err = (last_err or "") + f"\n[LO#2] exit={e2.returncode} STDOUT={str(e2.stdout)[:1000]} STDERR={str(e2.stderr)[:1000]}"

            if last_err is None:
                candidates = ["documento.pdf", Path(in_path).with_suffix(".pdf").name]
                for name in candidates:
                    pdf_path = os.path.join(out_dir, name)
                    if os.path.exists(pdf_path):
                        with open(pdf_path, "rb") as f:
                            return io.BytesIO(f.read())
                for name in os.listdir(out_dir):
                    if name.lower().endswith(".pdf"):
                        with open(os.path.join(out_dir, name), "rb") as f:
                            return io.BytesIO(f.read())
                last_err = "Conversão LO executou, mas o PDF não foi encontrado."

        try:
            docx2pdf_convert(in_path, out_dir)
            for name in ("documento.pdf", Path(in_path).with_suffix(".pdf").name):
                pdf_path = os.path.join(out_dir, name)
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        return io.BytesIO(f.read())
            for name in os.listdir(out_dir):
                if name.lower().endswith(".pdf"):
                    with open(os.path.join(out_dir, name), "rb") as f:
                        return io.BytesIO(f.read())
            raise RuntimeError("docx2pdf não gerou o PDF.")
        except Exception as e:
            hint = f"\nDetalhes LO: {last_err}" if last_err else ""
            logger.error("Falha na conversão DOCX→PDF. %s", hint)
            raise RuntimeError(
                "Falha na conversão DOCX→PDF. Ajuste o LibreOffice (SOFFICE_PATH) "
                "ou habilite o Microsoft Word (docx2pdf)."
            ) from e

def marca_agora():
    """Retorna timestamp (string) no fuso -03:00 (HH:MM:SS)."""
    return (datetime.utcnow() - timedelta(hours=3)).strftime("%Y-%m-%d %H:%M:%S")

def _safe_filename(name, ext):
    """Gera nome de arquivo seguro com extensão ext."""
    base = re.sub(r'[^\w\s.-]+', '', (name or 'documento')).strip() or 'documento'
    base = re.sub(r'\s+', ' ', base).replace(' ', '_')
    return f"{base}.{ext}"

def _load_componentes_e_variaveis(template_id):
    """Carrega lista de componentes, valores de variáveis e nome pelo id do template."""
    with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
        cur.execute("""
            SELECT nome_arquivo, template_html, variaveis_valores
              FROM galeria_juridico
             WHERE id_template = %s
        """, (template_id,))
        row = cur.fetchone()
        if not row:
            return None, None, None

        htmls = []
        try:
            maybe_list = json.loads(row["template_html"] or "[]")
            if isinstance(maybe_list, list):
                htmls = [c.get("template_html", "") for c in maybe_list]
            else:
                htmls = [row["template_html"] or ""]
        except Exception:
            htmls = [row["template_html"] or ""]

        vals = {}
        if row.get("variaveis_valores"):
            try:
                for item in json.loads(row["variaveis_valores"]):
                    n, v = item.get("nome"), item.get("valor")
                    if n:
                        vals[n] = v
            except Exception:
                pass

        return htmls, vals, row["nome_arquivo"]

def _apply_variables_to_html(raw_html, values_dict):
    """Aplica os valores de variáveis ao HTML (tanto spans data-variable quanto [VAR])."""
    html_out = raw_html or ""

    def repl_span(m):
        var_name = m.group(1)
        val = values_dict.get(var_name) if values_dict else None
        if val is None or str(val).strip() == "":
            return f'[{var_name}]'
        return html.escape(str(val))
    html_out = re.sub(
        r'<span[^>]*\sdata-variable="([^"]+)"[^>]*>.*?</span>',
        repl_span,
        html_out,
        flags=re.IGNORECASE | re.DOTALL
    )

    def repl_plain(m):
        var_name = m.group(1)
        val = values_dict.get(var_name) if values_dict else None
        return str(val) if (val is not None and str(val).strip() != "") else f'[{var_name}]'
    html_out = re.sub(r'\[([^\]]+)\]', repl_plain, html_out)

    return html_out

def _merge_componentes_com_variaveis(html_list, values_dict):
    """Mescla componentes aplicando variáveis e retorna HTML final."""
    parts = []
    for h in html_list or []:
        parts.append(_apply_variables_to_html(h, values_dict))
    return "".join(parts)

def _extract_pages_text_from_pdf(pdf_io, normalizar=True, manter_quebras=True):
    """Extrai lista de textos por página a partir de um PDF (BytesIO)."""
    pdf_io.seek(0)
    doc = fitz.open(stream=pdf_io.getvalue(), filetype="pdf")
    paginas = []
    for pno in range(doc.page_count):
        page = doc.load_page(pno)
        txt = page.get_text("text") or ""
        if normalizar:
            txt = txt.replace('\u00ad', '').replace('\xa0', ' ')
            if manter_quebras:
                linhas = []
                for ln in txt.splitlines():
                    ln = re.sub(r'\s+', ' ', ln).strip()
                    if ln:
                        linhas.append(ln)
                txt = "\n".join(linhas)
            else:
                txt = re.sub(r'\s+', ' ', txt).strip()
        paginas.append(txt)
    doc.close()
    return paginas

# === FLASK ===
app = Flask(__name__)
CORS(app)

@app.before_request
def _log_request_in():
    """Loga entrada de toda requisição (método, rota, IP, query, chaves do JSON)."""
    g._start_time = time.time()
    payload_keys = None
    if request.is_json:
        try:
            data = request.get_json(silent=True)
            if isinstance(data, dict):
                payload_keys = sorted(list(data.keys()))
            elif isinstance(data, list):
                payload_keys = f"list[{len(data)}]"
            else:
                payload_keys = str(type(data))
        except Exception:
            payload_keys = "unreadable-json"
    logger.info(
        "IN %s %s | ip=%s | args=%s | json_keys=%s",
        request.method, request.path, request.remote_addr,
        dict(request.args), payload_keys
    )

@app.after_request
def _log_request_out(response):
    """Loga saída de toda requisição (status e duração)."""
    try:
        dur_ms = (time.time() - getattr(g, "_start_time", time.time())) * 1000.0
        logger.info(
            "OUT %s %s | status=%s | dur_ms=%.1f",
            request.method, request.path, response.status_code, dur_ms
        )
    except Exception:
        logger.exception("after_request: falha ao logar saída")
    return response

@app.errorhandler(Exception)
def _handle_unexpected_error(e):
    """Captura exceções não tratadas, loga como ERROR e retorna JSON 500."""
    logger.exception("Exceção não tratada: %s", e)
    return jsonify(status="erro", mensagem="Erro interno no servidor."), 500

# === ENDPOINTS ===
@app.route("/biblioteca", methods=["GET"])
def biblioteca_unificada():
    sql_ativos = """
        SELECT id_template, nome_arquivo, usuario, data_criacao, caminho, alteracao, type, status
          FROM galeria_juridico
         WHERE type IN ('template','individual')
           AND COALESCE(status,'') NOT IN ('baixado','alterado')   -- <<<<<<<<<< ALTERAÇÃO
         ORDER BY id_template DESC
    """
    sql_baixados = """
        SELECT id_template, nome_arquivo, usuario, data_criacao, caminho, type
          FROM galeria_juridico
         WHERE status = 'baixado'
         ORDER BY data_criacao DESC, id_template DESC
    """
    sql_alt = """
        SELECT id_template, nome_arquivo, alteracao, usuario,
               COALESCE(data_alteracao, data_criacao) AS data_alteracao,
               caminho, type
          FROM galeria_juridico
         WHERE status='alterado'
           AND COALESCE(alteracao,'') IN ('alterado','deletado')
         ORDER BY data_alteracao DESC, id_template DESC
    """
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute(sql_ativos)
            rows_ativos = cur.fetchall()
            cur.execute(sql_baixados)
            rows_bx = cur.fetchall()
            cur.execute(sql_alt)
            rows_alt = cur.fetchall()

        def fmt_dt(v):
            return v.strftime("%d/%m/%Y %H:%M:%S") if isinstance(v, datetime) else (v or "")

        templates, individuais = [], []
        for r in rows_ativos:
            info_user = get_user_by_id(r["usuario"]) if r.get("usuario") else None
            item = {
                "id_template": r["id_template"],
                "titulo":      r["nome_arquivo"],
                "usuario":     info_user["NOME"].title() if info_user else "—",
                "data_criacao": fmt_dt(r["data_criacao"]),
                "caminho":     r["caminho"],
                "alteracao":   r["alteracao"],
                "tipo":        r["type"],
            }
            (templates if r["type"] == "template" else individuais).append(item)

        baixados = []
        for r in rows_bx:
            info_user = get_user_by_id(r["usuario"]) if r.get("usuario") else None
            baixados.append({
                "id_template":  r["id_template"],
                "titulo":       r["nome_arquivo"],
                "usuario":      info_user["NOME"].title() if info_user else "—",
                "data_download": fmt_dt(r["data_criacao"]),
                "caminho":      r["caminho"],
                "tipo":         r["type"],
            })

        pastas_excluidas = {
            f"{r['caminho'] or '/'}{r['nome_arquivo']}/"
            for r in rows_alt if r["type"] == "pasta"
        }
        def dentro_de_pasta_excluida(reg):
            abs_path = f"{reg['caminho'] or '/'}{reg['nome_arquivo']}"
            return any(abs_path.startswith(pref) for pref in pastas_excluidas)

        alteracoes = []
        for r in rows_alt:
            if dentro_de_pasta_excluida(r):
                continue
            info_user = get_user_by_id(r["usuario"]) if r.get("usuario") else None
            alteracoes.append({
                "id_template":   r["id_template"],
                "nome_arquivo":  r["nome_arquivo"],
                "alteracao":     r["alteracao"],
                "usuario":       info_user["NOME"].title() if info_user else "—",
                "data_alteracao": fmt_dt(r["data_alteracao"]),
                "caminho":       r["caminho"],
                "type":          r["type"],
            })

        return jsonify(
            status="sucesso",
            mensagem="Listagem da biblioteca obtida com sucesso.",
            templates=templates,
            individuais=individuais,
            baixados=baixados,
            alteracoes=alteracoes,
        )
    except Exception:
        logger.exception("/biblioteca: falha ao listar biblioteca")
        return jsonify(status="erro", mensagem="Falha ao listar biblioteca"), 500

@app.route("/nova_pasta", methods=["POST"])
def nova_pasta():
    data = request.get_json(force=True)
    caminho = data.get("caminho") or "/"
    usuario = (data.get("usuario") or data.get("id_usuario") or "").strip() or None
    try:
        with get_connection() as cnx:
            nome_final = nome_unico(cnx, caminho, data["nome_arquivo"])
            sql = """
                INSERT INTO galeria_juridico
                       (nome_arquivo, status, alteracao,
                        id_origem,    usuario,
                        data_criacao, caminho, type)
                VALUES (%s, %s, %s, %s, %s, NOW(), %s, 'pasta')
            """
            last_id = executar(sql, (nome_final, data.get("status",""), None,
                                     data.get("id_origem"), usuario, caminho))
        return jsonify(status="sucesso", mensagem="Pasta criada com sucesso.", id_template=last_id, nome=nome_final)
    except Exception:
        logger.exception("/nova_pasta: falha ao criar pasta")
        return jsonify(status="erro", mensagem="Falha ao criar pasta"), 500

word_mimes = {
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

@app.route("/novo_template", methods=["POST"])
def novo_template():
    """Cadastra um novo template a partir de um arquivo .doc/.docx."""
    if "arquivo" not in request.files:
        return jsonify(status="erro", mensagem="Arquivo ausente"), 400
    file = request.files["arquivo"]
    if (file.mimetype not in word_mimes or not file.filename.lower().endswith((".doc", ".docx"))):
        return jsonify(status="erro", mensagem="Tipo de arquivo inválido"), 400

    blob = file.read()
    form = request.form
    caminho = form.get("caminho") or "/"
    # <- ACEITA usuario OU id_usuario do FormData
    usuario = (form.get("usuario") or form.get("id_usuario") or "").strip() or None
    nome_base = (form.get("nome_arquivo") or file.filename).rsplit(".", 1)[0]

    try:
        html_paginado = docx_to_html(blob)
        lista_componentes = json.dumps([{"nome_arquivo": nome_base, "template_html": html_paginado}], ensure_ascii=False)
        with get_connection() as cnx:
            nome_final = nome_unico(cnx, caminho, nome_base)
            sql = """
                INSERT INTO galeria_juridico
                       (nome_arquivo, status, alteracao,
                        id_origem,    usuario,
                        data_criacao, caminho, type,
                        template,     template_html)
                VALUES (%s, %s, %s,
                        %s, %s,
                        NOW(), %s, 'template',
                        %s, %s)
            """
            last_id = executar(sql, (
                nome_final,
                form.get("status", ""),
                None,
                form.get("id_origem"),
                usuario,           # <- gravando corretamente
                caminho,
                blob,
                lista_componentes
            ))
        return jsonify(status="sucesso", mensagem="Template cadastrado com sucesso.", id_template=last_id, nome=nome_final)
    except Exception:
        logger.exception("/novo_template: falha ao processar arquivo")
        return jsonify(status="erro", mensagem="Não foi possível processar o arquivo"), 500

@app.route("/listar_galeria", methods=["GET"])
def listar_galeria():
    """Lista todos os registros não deletados da galeria (ordem decrescente por id)."""
    sql = """
        SELECT id_template,
               nome_arquivo,
               status,
               alteracao,
               id_origem,
               usuario,
               data_criacao,
               data_alteracao,
               caminho,
               type
          FROM galeria_juridico
         WHERE NOT (
                   status = 'deletado'
                   OR (status = 'alterado' AND COALESCE(alteracao,'') = 'deletado')
               )
         ORDER BY id_template DESC
    """
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute(sql)
            rows = cur.fetchall()
        for r in rows:
            for campo in ("data_criacao", "data_alteracao"):
                if isinstance(r[campo], datetime):
                    r[campo] = r[campo].strftime("%Y-%m-%d %H:%M:%S")
        return jsonify(status="sucesso", mensagem="Listagem da galeria obtida com sucesso.", registros=rows)
    except Exception:
        logger.exception("/listar_galeria: falha na consulta")
        return jsonify(status="erro", mensagem="Falha ao listar registros"), 500

@app.route("/alterar_status", methods=["POST"])
def alterar_status():
    """Altera status/alteração de arquivo/pasta; em pasta, aplica em cascata por caminho.
       Se o documento estiver em edição por outro usuário, bloqueia.
       Se o lock for do mesmo usuário (em_uso_por == usuario), permite."""
    data = request.get_json(force=True)
    item_id = data.get("id_template")
    novo_status = (data.get("status") or "").strip().lower()
    alteracao = data.get("alteracao") or None
    usuario_req = (data.get("usuario") or "").strip()

    if not item_id:
        return jsonify(status="erro", mensagem="Parâmetros ausentes")
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("""
                SELECT nome_arquivo, caminho, type, em_uso, em_uso_modo, em_uso_por
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (item_id,))
            row = cur.fetchone()
            if not row:
                return jsonify(status="sucesso", mensagem="Item não encontrado; nada a alterar.")

            # bloqueia apenas se em edição por OUTRO usuário
            if row["em_uso"] and (row.get("em_uso_modo") or "edicao") != "leitura" and str(row.get("em_uso_por") or "") != str(usuario_req or ""):
                return jsonify(status="erro", mensagem="Documento já está sendo editado")

            if row["type"] == "pasta":
                pasta_abs = f"{row['caminho'] or '/'}{row['nome_arquivo']}/"
                cur.execute("""
                    UPDATE galeria_juridico
                       SET status         = %s,
                           alteracao      = %s,
                           data_alteracao = NOW()
                     WHERE id_template   = %s
                        OR caminho LIKE %s
                """, (novo_status, alteracao, item_id, f"{pasta_abs}%"))
            else:
                cur.execute("""
                    UPDATE galeria_juridico
                       SET status         = %s,
                           alteracao      = %s,
                           data_alteracao = NOW()
                     WHERE id_template   = %s
                """, (novo_status, alteracao, item_id))
            cnx.commit()
        return jsonify(status="sucesso", mensagem="Status atualizado com sucesso.")
    except Exception:
        logger.exception("/alterar_status: falha no update")
        return jsonify(status="erro", mensagem="Falha ao alterar status"), 500

@app.route("/reverter_alteracao", methods=["POST"])
def reverter_alteracao():
    data = request.get_json(force=True) or {}
    id_alt = data.get("id_alteracao") or data.get("id")
    usuario_req = (data.get("usuario") or data.get("id_usuario") or "").strip()

    if not id_alt:
        return jsonify(status="erro", mensagem="Parâmetros ausentes (id_alteracao)"), 400

    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            # 1) variação
            cur.execute("""
                SELECT id_template, id_origem, status, alteracao,
                       nome_arquivo, caminho, type
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (id_alt,))
            var = cur.fetchone()
            if not var:
                return jsonify(status="erro", mensagem="Variação não encontrada"), 404
            if (var.get("status") or "").lower() != "alterado" or (var.get("alteracao") or "").lower() != "alterado":
                return jsonify(status="erro", mensagem="O item informado não é uma variação de alteração"), 400

            id_pai = var.get("id_origem")
            pai_lock = None

            # 2) tenta ler o pai direto
            if id_pai:
                cur.execute("""
                    SELECT id_template, em_uso, em_uso_modo, em_uso_por
                      FROM galeria_juridico
                     WHERE id_template = %s
                """, (id_pai,))
                pai_lock = cur.fetchone()

            # 3) fallback: pai já foi removido – tenta achar o "principal" atual compatível
            if not pai_lock:
                cur.execute("""
                    SELECT id_template, em_uso, em_uso_modo, em_uso_por
                      FROM galeria_juridico
                     WHERE nome_arquivo = %s
                       AND caminho      = %s
                       AND type         = %s
                       AND COALESCE(status,'') NOT IN ('baixado','alterado')
                     ORDER BY id_template DESC
                     LIMIT 1
                """, (var["nome_arquivo"], var["caminho"], var["type"]))
                pai_lock = cur.fetchone()
                id_pai = pai_lock["id_template"] if pai_lock else None

            # 4) bloqueio só se alguém estiver EDITANDO um possível pai atual
            if pai_lock and pai_lock.get("em_uso") and (pai_lock.get("em_uso_modo") or "edicao") == "edicao":
                if str(pai_lock.get("em_uso_por") or "") != str(usuario_req or ""):
                    return jsonify(status="erro", mensagem="Template pai em edição por outro usuário"), 409

            agora = marca_agora()

            # 5) promove variação
            cur.execute("""
                UPDATE galeria_juridico
                   SET status = '',
                       alteracao = NULL,
                       id_origem = NULL,
                       em_uso = FALSE,
                       em_uso_modo = NULL,
                       em_uso_por = NULL,
                       data_alteracao = %s
                 WHERE id_template = %s
            """, (agora, id_alt,))

            # 6) remove pai (se houver)
            if id_pai:
                cur.execute("DELETE FROM galeria_juridico WHERE id_template = %s", (id_pai,))

            cnx.commit()

        return jsonify(
            status="sucesso",
            mensagem="Variação promovida e template original removido.",
            id_promovido=id_alt
        )
    except Exception:
        logger.exception("/reverter_alteracao: falha")
        return jsonify(status="erro", mensagem="Falha ao reverter alteração"), 500

@app.route("/renomear", methods=["POST"])
def renomear():
    """Renomeia arquivo ou pasta, mantendo unicidade no caminho.
       Se estiver em uso por OUTRO usuário, bloqueia. Se for o mesmo, permite."""
    data = request.get_json(force=True)
    item_id   = data.get("id_template")
    novo_nome = (data.get("novo_nome") or "").strip()
    usuario_req = (data.get("usuario") or "").strip()

    if not item_id or not novo_nome:
        return jsonify(status="erro", mensagem="Parâmetros ausentes")

    novo_nome_base = novo_nome.rsplit(".", 1)[0]
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("""
                SELECT caminho, alteracao, type, nome_arquivo, em_uso, em_uso_modo, em_uso_por
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (item_id,))
            row = cur.fetchone()
            if not row:
                return jsonify(status="erro", mensagem="ID não encontrado")
            if row["alteracao"] == "deletado":
                return jsonify(status="erro", mensagem="Item deletado")

            # Bloqueia apenas se estiver em uso por OUTRO usuário
            if row["em_uso"] and str(row.get("em_uso_por") or "") != str(usuario_req or ""):
                return jsonify(status="erro", mensagem="Documento já está sendo editado")

            caminho_atual = row["caminho"] or "/"
            nome_final = nome_unico(cnx, caminho_atual, novo_nome_base)

            if row["type"] == "pasta":
                antigo_abs = f"{caminho_atual}{row['nome_arquivo']}/"
                novo_abs   = f"{caminho_atual}{nome_final}/"
                cur.execute("UPDATE galeria_juridico SET nome_arquivo=%s WHERE id_template=%s",
                            (nome_final, item_id))
                cur.execute("""
                    UPDATE galeria_juridico
                       SET caminho = REPLACE(caminho, %s, %s)
                     WHERE caminho LIKE %s
                """, (antigo_abs, novo_abs, f"{antigo_abs}%"))
            else:
                cur.execute("UPDATE galeria_juridico SET nome_arquivo=%s WHERE id_template=%s",
                            (nome_final, item_id))
            cnx.commit()
        return jsonify(status="sucesso", mensagem="Item renomeado com sucesso.", nome=nome_final)
    except Exception:
        logger.exception("/renomear: falha no rename")
        return jsonify(status="erro", mensagem="Falha ao renomear"), 500

@app.route("/mover", methods=["POST"])
def mover():
    """Move item para outro caminho garantindo nome único no destino."""
    data = request.get_json(force=True)
    item_id = int(data["id_template"])
    novo_caminho = data.get("novo_caminho") or "/"
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("SELECT nome_arquivo FROM galeria_juridico WHERE id_template=%s",(item_id,))
            row = cur.fetchone()
            if not row:
                return jsonify(status="erro", mensagem="ID não encontrado")
            nome_final = nome_unico(cnx, novo_caminho, row["nome_arquivo"])
            cur.execute(
                "UPDATE galeria_juridico SET caminho=%s, nome_arquivo=%s WHERE id_template=%s",
                (novo_caminho, nome_final, item_id)
            )
            cnx.commit()
        return jsonify(status="sucesso", mensagem="Item movido com sucesso.", nome=nome_final)
    except Exception:
        logger.exception("/mover: falha ao mover item")
        return jsonify(status="erro", mensagem="Falha ao mover"), 500

@app.route("/duplicar_template", methods=["POST"])
def duplicar_template():
    """Duplica template/individual para novo registro (tipo_destino: 'template' ou 'individual')."""
    data = request.get_json(force=True) or {}
    src_id = data.get("id_template")
    tipo_dest = (data.get("tipo_destino") or "template").lower()
    if tipo_dest not in ("template", "individual"):
        return jsonify(status="erro", mensagem="tipo_destino inválido")
    if not src_id:
        return jsonify(status="erro", mensagem="ID ausente")

    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("""
                SELECT nome_arquivo, status, usuario, caminho,
                       template, template_html
                  FROM galeria_juridico
                 WHERE id_template = %s AND type IN ('template','individual')
            """, (src_id,))
            tpl = cur.fetchone()
            if not tpl:
                return jsonify(status="erro", mensagem="Template não encontrado")

            novo_nome = nome_unico(cnx, tpl["caminho"] or "/", tpl["nome_arquivo"])
            new_id = executar("""
                INSERT INTO galeria_juridico
                       (nome_arquivo, status, alteracao,
                        id_origem,    usuario,
                        data_criacao, caminho, type,
                        template,     template_html)
                VALUES (%s, %s, %s,
                        %s, %s,
                        NOW(), %s, %s,
                        %s, %s)
            """, (
                novo_nome,
                tpl["status"] or "",
                None,
                src_id,
                tpl["usuario"],
                tpl["caminho"],
                tipo_dest,
                tpl["template"],
                tpl["template_html"],
            ))
        return jsonify(status="sucesso", mensagem="Documento duplicado com sucesso.", id_template=new_id, nome=novo_nome, tipo=tipo_dest)
    except Exception:
        logger.exception("/duplicar_template: falha ao duplicar")
        return jsonify(status="erro", mensagem="Falha ao duplicar"), 500

@app.route("/registrar_download", methods=["POST"])
def registrar_download():
    """Cria um snapshot 'baixado' do template/individual (status='baixado')."""
    data = request.get_json(force=True) or {}
    src_id  = data.get("id_template")
    usuario = data.get("usuario")
    if not src_id:
        return jsonify(status="erro", mensagem="id_template ausente"), 400
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("""
                SELECT id_template, nome_arquivo, usuario AS src_user, caminho, type,
                       template, template_html, variaveis_valores
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (src_id,))
            base = cur.fetchone()
            if not base:
                return jsonify(status="erro", mensagem="Template base não encontrado"), 404

            user_final = usuario or base.get("src_user")
            if not user_final:
                user_final = None

            new_id = executar("""
                INSERT INTO galeria_juridico
                       (nome_arquivo, status, alteracao,
                        id_origem,    usuario,
                        data_criacao, caminho, type,
                        template,     template_html, variaveis_valores, em_uso)
                VALUES (%s, 'baixado', NULL,
                        %s, %s,
                        NOW(), %s, %s,
                        %s, %s, %s, 0)
            """, (
                base["nome_arquivo"],
                base["id_template"],
                user_final,
                base["caminho"],
                base["type"],
                base["template"],
                base["template_html"],
                base["variaveis_valores"]
            ))
        return jsonify(status="sucesso", mensagem="Download registrado com sucesso.", id_template=new_id)
    except Exception:
        logger.exception("/registrar_download: falha ao registrar")
        return jsonify(status="erro", mensagem="Falha ao registrar download"), 500

@app.route("/reverter_download", methods=["POST"])
def reverter_download():
    """
    Promove o snapshot 'baixado' (id_baixado) a template principal:
      - variação: status='', alteracao=NULL, id_origem=NULL, limpa locks
      - remove o pai (id_origem) se existir
      - se o pai não existir mais, tenta achar o 'principal' atual compatível
        (mesmo nome_arquivo, caminho e type, com status not in ('baixado','alterado'))
      - bloqueia apenas se o pai (encontrado) estiver em edição por OUTRO usuário
    Retorna id_promovido (id do snapshot promovido).
    """
    data = request.get_json(force=True) or {}
    id_baixado = data.get("id_baixado") or data.get("id")
    usuario_req = (data.get("usuario") or data.get("id_usuario") or "").strip()

    if not id_baixado:
        return jsonify(status="erro", mensagem="Parâmetros ausentes (id_baixado)"), 400

    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            # 1) Carrega o snapshot
            cur.execute("""
                SELECT id_template, id_origem, status, nome_arquivo, caminho, type
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (id_baixado,))
            snap = cur.fetchone()
            if not snap:
                return jsonify(status="erro", mensagem="Snapshot não encontrado"), 404

            if (snap.get("status") or "").lower() != "baixado":
                return jsonify(status="erro", mensagem="O item informado não é um registro de download"), 400

            # 2) Tenta localizar o pai pelo id_origem
            id_pai = snap.get("id_origem")
            pai_lock = None
            if id_pai:
                cur.execute("""
                    SELECT id_template, em_uso, em_uso_modo, em_uso_por
                      FROM galeria_juridico
                     WHERE id_template = %s
                """, (id_pai,))
                pai_lock = cur.fetchone()

            # 3) Fallback: se o pai não existir mais, procura um "principal" compatível
            #    (mesmo nome/caminho/type e status que NÃO seja 'baixado' nem 'alterado')
            if not pai_lock:
                cur.execute("""
                    SELECT id_template, em_uso, em_uso_modo, em_uso_por
                      FROM galeria_juridico
                     WHERE nome_arquivo = %s
                       AND caminho      = %s
                       AND type         = %s
                       AND COALESCE(status,'') NOT IN ('baixado','alterado')
                     ORDER BY id_template DESC
                     LIMIT 1
                """, (snap["nome_arquivo"], snap["caminho"], snap["type"]))
                pai_lock = cur.fetchone()
                id_pai = pai_lock["id_template"] if pai_lock else None

            # 4) Se existe um pai "vivo" e ele está em edição por OUTRO usuário, bloqueia
            if pai_lock and pai_lock.get("em_uso") and (pai_lock.get("em_uso_modo") or "edicao") == "edicao":
                if str(pai_lock.get("em_uso_por") or "") != str(usuario_req or ""):
                    return jsonify(status="erro", mensagem="Template pai em edição por outro usuário"), 409

            agora = marca_agora()

            # 5) Promove o snapshot: torna-se o principal e limpa locks
            cur.execute("""
                UPDATE galeria_juridico
                   SET status         = '',
                       alteracao      = NULL,
                       id_origem      = NULL,
                       em_uso         = FALSE,
                       em_uso_modo    = NULL,
                       em_uso_por     = NULL,
                       data_alteracao = %s
                 WHERE id_template    = %s
            """, (agora, id_baixado,))

            # 6) Remove o pai (se ainda existir)
            if id_pai:
                cur.execute("DELETE FROM galeria_juridico WHERE id_template = %s", (id_pai,))

            cnx.commit()

        return jsonify(
            status="sucesso",
            mensagem="Snapshot promovido e template original removido.",
            id_promovido=id_baixado
        )
    except Exception:
        logger.exception("/reverter_download: falha ao reverter")
        return jsonify(status="erro", mensagem="Falha ao reverter download"), 500

@app.route("/delete_permanente", methods=["POST"])
def delete_permanente():
    data = request.get_json(force=True) or {}
    item_id = data.get("id_template")
    tipo_req = (data.get("tipo") or "").lower().strip()
    usuario_req = (data.get("usuario") or "").strip()
    if not item_id:
        return jsonify(status="erro", mensagem="Parâmetros ausentes"), 400
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("""
                SELECT id_template, type, nome_arquivo, caminho, em_uso, em_uso_modo, em_uso_por, status, alteracao
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (item_id,))
            raiz = cur.fetchone()
            if not raiz:
                return jsonify(status="erro", mensagem="ID não encontrado"), 404

            # Bloqueia apenas se em edição por OUTRO usuário
            if raiz["em_uso"] and (raiz.get("em_uso_modo") or "edicao") != "leitura" and str(raiz.get("em_uso_por") or "") != str(usuario_req or ""):
                return jsonify(status="erro", mensagem="Documento já está sendo editado"), 409

            tipo = (tipo_req or raiz["type"] or "").lower()
            ids_all = {raiz["id_template"]}

            if tipo == "pasta":
                abs_path = f"{raiz['caminho'] or '/'}{raiz['nome_arquivo']}/"
                cur.execute("""
                    SELECT id_template
                      FROM galeria_juridico
                     WHERE caminho LIKE %s
                       AND status='alterado'
                       AND COALESCE(alteracao,'')='deletado'
                """, (f"{abs_path}%",))
                ids_all |= {r["id_template"] for r in cur.fetchall()}

            ids_del = sorted(ids_all)

            if ids_del:
                placeholders = ",".join(["%s"] * len(ids_del))

                # Itens em edição por OUTROS usuários (ignora locks do mesmo usuário)
                cur.execute(
                    f"""SELECT id_template, nome_arquivo
                           FROM galeria_juridico
                          WHERE em_uso = 1
                            AND (em_uso_modo IS NULL OR em_uso_modo <> 'leitura')
                            AND (em_uso_por IS NULL OR em_uso_por <> %s)
                            AND id_template IN ({placeholders})""",
                    (usuario_req, *ids_del)
                )
                locked = cur.fetchall()
                if locked:
                    nomes = ", ".join(f"#{r['id_template']} {r['nome_arquivo']}" for r in locked)
                    return jsonify(status="erro", mensagem=f"Um ou mais itens estão em edição: {nomes}"), 409

                # Libera locks em leitura (se houver)
                cur.execute(
                    f"""UPDATE galeria_juridico
                           SET em_uso = FALSE, em_uso_modo = NULL, em_uso_por = NULL
                         WHERE em_uso = TRUE
                           AND em_uso_modo = 'leitura'
                           AND id_template IN ({placeholders})""",
                    tuple(ids_del)
                )

                cur.execute(f"DELETE FROM galeria_juridico WHERE id_template IN ({placeholders})", tuple(ids_del))
                cnx.commit()

        return jsonify(status="sucesso", mensagem="Itens removidos permanentemente com sucesso.", removidos=len(ids_del))
    except Exception:
        logger.exception("/delete_permanente: falha ao excluir")
        return jsonify(status="erro", mensagem="Falha na exclusão"), 500

@app.route("/get_template", methods=["GET"])
def get_template():
	"""Abre template para edição/leitura, aplicando lock adequado e permitindo reentrância do mesmo usuário."""
	template_id = request.args.get("id", type=int)
	modo = (request.args.get("modo") or "edicao").strip().lower()
	# Aceita 'usuario' OU 'id_usuario' vindos do front
	usuario_req = (request.args.get("usuario") or request.args.get("id_usuario") or "").strip() or None

	if not template_id:
		return jsonify(status="erro", mensagem="ID faltando")

	try:
		with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
			# Lê o estado atual do lock
			cur.execute("""
				SELECT em_uso, em_uso_modo, em_uso_por
				  FROM galeria_juridico
				 WHERE id_template = %s
			""", (template_id,))
			lock = cur.fetchone() or {}
			em_uso  = bool(lock.get("em_uso"))
			em_modo = (lock.get("em_uso_modo") or "edicao")
			em_por  = (lock.get("em_uso_por") or None)
			same_user = (usuario_req is not None and str(em_por) == str(usuario_req))

			if modo == "leitura":
				# Se está em edição por outro usuário, bloqueia
				if em_uso and em_modo == "edicao" and not same_user:
					return jsonify(status="erro", mensagem="Documento já está sendo editado")
				# Se está em edição por este mesmo usuário, apenas renova timestamp
				if em_uso and em_modo == "edicao" and same_user:
					cur.execute("UPDATE galeria_juridico SET data_alteracao = NOW() WHERE id_template = %s", (template_id,))
				else:
					# Seta lock de leitura
					cur.execute("""
						UPDATE galeria_juridico
						   SET em_uso = TRUE,
							   em_uso_modo = 'leitura',
							   em_uso_por  = %s,
							   data_alteracao = NOW()
						 WHERE id_template = %s
					""", (usuario_req, template_id))
				cnx.commit()
			else:
				# modo 'edicao'
				# Se está em edição por outro usuário, bloqueia
				if em_uso and em_modo == "edicao" and not same_user:
					return jsonify(status="erro", mensagem="Documento já está sendo editado")
				# Permite reentrância do mesmo usuário e também tomar do 'leitura'
				cur.execute("""
					UPDATE galeria_juridico
					   SET em_uso = TRUE,
						   em_uso_modo = 'edicao',
						   em_uso_por  = %s,
						   data_alteracao = NOW()
					 WHERE id_template = %s
				""", (usuario_req, template_id))
				cnx.commit()

			# Carrega conteúdo do template
			cur.execute("""
				SELECT nome_arquivo, template, template_html, type, data_alteracao, variaveis_valores
				  FROM galeria_juridico
				 WHERE id_template = %s
			""", (template_id,))
			tpl = cur.fetchone()
			if not tpl:
				# Reverte lock caso não exista mais
				executar("UPDATE galeria_juridico SET em_uso = FALSE, em_uso_modo=NULL, em_uso_por=NULL WHERE id_template = %s", (template_id,))
				return jsonify(status="erro", mensagem="Template não encontrado")

			if not tpl["template_html"]:
				html_conv = docx_to_html(tpl["template"])
				cur.execute(
					"UPDATE galeria_juridico SET template_html = %s, data_alteracao = NOW() WHERE id_template = %s",
					(html_conv, template_id))
				cnx.commit()
				tpl["template_html"] = html_conv

			try:
				componentes = json.loads(tpl["template_html"])
				assert isinstance(componentes, list)
			except Exception:
				componentes = [{"nome_arquivo": tpl["nome_arquivo"], "template_html": tpl["template_html"]}]

			valores_salvos = []
			if tpl.get("variaveis_valores"):
				try:
					valores_salvos = json.loads(tpl["variaveis_valores"])
				except Exception:
					valores_salvos = []

			variaveis_encontradas = set()
			for comp in componentes:
				soup = BeautifulSoup(comp["template_html"], 'html.parser')
				texto_puro = soup.get_text()
				matches = re.findall(r'\[([^\]]+)\]', texto_puro)
				for v in matches:
					variaveis_encontradas.add(v)

			variaveis = list(variaveis_encontradas)
			hora_iso = tpl["data_alteracao"].isoformat(sep=" ", timespec="seconds") if tpl["data_alteracao"] else marca_agora()

		return jsonify(
			status="sucesso",
			mensagem="Template aberto com sucesso.",
			nome_arquivo=tpl["nome_arquivo"],
			tipo=tpl["type"],
			hora_servidor=hora_iso,
			componentes=componentes,
			variaveis=variaveis,
			variaveis_valores=valores_salvos,
			total_variaveis=len(variaveis)
		)
	except Exception as exc:
		try:
			executar("UPDATE galeria_juridico SET em_uso = FALSE, em_uso_modo=NULL, em_uso_por=NULL WHERE id_template = %s", (template_id,))
		except Exception:
			pass
		logger.exception("/get_template: falha ao abrir template")
		return jsonify(status="erro", mensagem=str(exc))

@app.route("/listar_variaveis", methods=["GET"])
def listar_variaveis():
    """Lista variáveis cadastradas (tabela variaveis_templates)."""
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("SELECT * FROM variaveis_templates ORDER BY nome_variavel")
            variaveis = cur.fetchall()
        return jsonify(status="sucesso", mensagem="Variáveis listadas com sucesso.", variaveis=variaveis)
    except Exception:
        logger.exception("/listar_variaveis: falha ao listar")
        return jsonify(status="erro", mensagem="Falha ao listar variáveis"), 500

@app.route("/nova_variavel", methods=["POST"])
def nova_variavel():
    """Cadastra nova variável (nome_variavel)."""
    data = request.get_json(force=True)
    nome_variavel = data.get("nome_variavel")
    if not nome_variavel:
        return jsonify(status="erro", mensagem="Nome da variável é obrigatório")
    try:
        with get_connection() as cnx, cnx.cursor() as cur:
            cur.execute("""
                INSERT INTO variaveis_templates (nome_variavel) VALUES (%s)
            """, (nome_variavel,))
            cnx.commit()
            last_id = cur.lastrowid
        return jsonify(status="sucesso", mensagem="Variável cadastrada com sucesso.", id_variavel=last_id)
    except mysql.connector.IntegrityError:
        logger.warning("/nova_variavel: variável já existe (%s)", nome_variavel)
        return jsonify(status="erro", mensagem="Variável já existe")
    except Exception:
        logger.exception("/nova_variavel: erro interno")
        return jsonify(status="erro", mensagem="Erro interno ao cadastrar variável")

@app.route("/delete_variavel", methods=["DELETE"])
def remover_variavel():
    """Remove variável pelo nome (nome_variavel)."""
    data = request.get_json(force=True)
    nome_variavel = data.get("nome_variavel")
    try:
        with get_connection() as cnx, cnx.cursor() as cur:
            cur.execute("DELETE FROM variaveis_templates WHERE nome_variavel = %s", (nome_variavel,))
            cnx.commit()
        return jsonify(status="sucesso", mensagem="Variável removida com sucesso.")
    except Exception:
        logger.exception("/delete_variavel: falha ao remover")
        return jsonify(status="erro", mensagem="Falha ao remover variável"), 500

@app.route("/renomear_componente", methods=["POST"])
def renomear_componente():
    """Renomeia o nome_arquivo de um componente (pelo índice) dentro de um template."""
    data = request.get_json(force=True) or {}
    id_template = data.get("id_template")
    novo_nome = (data.get("novo_nome") or "").strip()
    indice = data.get("indice")
    if not id_template or not novo_nome or indice is None:
        return jsonify(status="erro", mensagem="Parâmetros ausentes")
    try:
        indice = int(indice)
        if indice < 0:
            raise ValueError
    except Exception:
        return jsonify(status="erro", mensagem="Índice inválido")

    agora = marca_agora()
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("SELECT template_html FROM galeria_juridico WHERE id_template=%s",(id_template,))
            row = cur.fetchone()
            if not row:
                return jsonify(status="erro", mensagem="ID não encontrado")

            try:
                lista = json.loads(row["template_html"] or "[]")
                assert isinstance(lista, list)
            except Exception:
                return jsonify(status="erro", mensagem="Formato de template_html inválido")

            if indice >= len(lista):
                return jsonify(status="erro", mensagem="Índice fora de alcance")

            lista[indice]["nome_arquivo"] = novo_nome
            cur.execute("""
                UPDATE galeria_juridico
                   SET template_html  = %s,
                       data_alteracao = %s
                 WHERE id_template    = %s
            """, (json.dumps(lista, ensure_ascii=False), agora, id_template))
            cnx.commit()
        return jsonify(status="sucesso", mensagem="Componente renomeado com sucesso.", data_alteracao=agora)
    except Exception:
        logger.exception("/renomear_componente: falha ao renomear")
        return jsonify(status="erro", mensagem="Falha ao renomear componente"), 500

@app.route("/salvar_documento", methods=["POST"])
def salvar_documento():
    """
    Atualiza SEMPRE o template pai com o conteúdo novo.
    Se 'fechar' == True e houve mudança em relação ao que está no banco,
    cria UMA variação em Alterações contendo o CONTEÚDO ANTERIOR (pré-salvar),
    com status='alterado' e alteracao='alterado'. Não altera o status do pai.
    """
    data = request.get_json(force=True) or {}

    id_template       = data.get("id_template")
    componentes       = data.get("componentes")
    liberar           = data.get("liberar", False)
    fechar_flag       = bool(data.get("fechar", False))  # <<<<<<<<<<<<<< ALTERAÇÃO
    variaveis_valores = data.get("variaveis_valores")
    usuario_req       = (data.get("usuario") or data.get("id_usuario") or "").strip() or None

    if not id_template or componentes is None:
        return jsonify(status="erro", mensagem="Parâmetros ausentes: id_template e componentes são obrigatórios")
    if not isinstance(componentes, list):
        return jsonify(status="erro", mensagem="O campo 'componentes' deve ser uma lista")
    if not isinstance(liberar, bool):
        return jsonify(status="erro", mensagem="O campo 'liberar' deve ser booleano")

    try:
        novo_template_html = json.dumps(componentes, ensure_ascii=False)
        valores_json = json.dumps(variaveis_valores, ensure_ascii=False) if variaveis_valores is not None else None
    except Exception:
        return jsonify(status="erro", mensagem="Dados inválidos para conversão em JSON")

    agora = marca_agora()
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            # lê estado atual
            cur.execute("""
                SELECT nome_arquivo, caminho, type, usuario,
                       template, template_html, variaveis_valores, data_alteracao
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (id_template,))
            row = cur.fetchone()
            if not row:
                return jsonify(status="erro", mensagem="ID do template não encontrado no banco de dados")

            old_template_html = row.get("template_html") or ""
            old_vals_json     = row.get("variaveis_valores")

            conteudo_mudou = (old_template_html != novo_template_html) or (old_vals_json != valores_json)

            # 1) sempre atualiza o pai
            cur.execute("""
                UPDATE galeria_juridico
                   SET template_html     = %s,
                       variaveis_valores = %s,
                       data_alteracao    = %s
                 WHERE id_template       = %s
            """, (novo_template_html, valores_json, agora, id_template))

            new_alt_id = None

            # 2) se fechando e houve mudança → cria variação com CONTEÚDO ANTERIOR
            if fechar_flag and conteudo_mudou:
                cur.execute("""
                    INSERT INTO galeria_juridico
                           (nome_arquivo, status, alteracao,
                            id_origem,    usuario,
                            data_criacao, caminho, type,
                            template,     template_html, variaveis_valores, em_uso)
                    VALUES (%s, 'alterado', 'alterado',
                            %s, %s,
                            NOW(), %s, %s,
                            %s, %s, %s, 0)
                """, (
                    row["nome_arquivo"],
                    id_template,
                    (usuario_req or row.get("usuario")),
                    (row.get("caminho") or "/"),
                    row.get("type"),
                    row.get("template"),
                    old_template_html,
                    old_vals_json
                ))
                new_alt_id = cur.lastrowid

            cnx.commit()

        return jsonify(
            status="sucesso",
            mensagem="Documento salvo com sucesso.",
            data_alteracao=agora,
            foi_alterado=conteudo_mudou,
            id_alteracao=new_alt_id
        )
    except Exception:
        logger.exception("/salvar_documento: falha ao salvar")
        return jsonify(status="erro", mensagem="Falha ao salvar documento"), 500

@app.route("/salvar_como_novo", methods=["POST"])
def salvar_como_novo():
    """Cria um novo documento na galeria a partir de um componente de outro template."""
    data = request.get_json(force=True) or {}
    id_base   = data.get("id_template_base")
    idx_comp  = data.get("idx_comp")
    html_comp = (data.get("html") or "").strip()
    nome_req  = (data.get("nome_arquivo") or "").strip()
    tipo_dest = (data.get("tipo_destino") or "individual").strip().lower()
    # <- ACEITA usuario OU id_usuario
    usuario_req = (data.get("usuario") or data.get("id_usuario") or "").strip() or None
    caminho_d = data.get("caminho_destino")
    fail_if_exists = bool(data.get("fail_if_exists", False))

    if tipo_dest not in ("template", "individual"):
        tipo_dest = "individual"
    if id_base is None or idx_comp is None or not html_comp:
        return jsonify(status="erro", mensagem="Parâmetros ausentes (id_template_base, idx_comp, html são obrigatórios)"), 400
    try:
        idx_comp = int(idx_comp)
    except Exception:
        return jsonify(status="erro", mensagem="Índice do componente inválido"), 400

    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("""
                SELECT nome_arquivo, caminho, usuario, template_html
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (id_base,))
            base = cur.fetchone()
            if not base:
                return jsonify(status="erro", mensagem="Template base não encontrado"), 404

            comp_nome = None
            try:
                comp_list = json.loads(base["template_html"] or "[]")
                if isinstance(comp_list, list) and 0 <= idx_comp < len(comp_list):
                    comp_nome = comp_list[idx_comp].get("nome_arquivo") or base["nome_arquivo"]
            except Exception:
                pass
            if not comp_nome:
                comp_nome = base["nome_arquivo"]

            nome_base_final = nome_req or comp_nome
            caminho_final   = caminho_d or (base.get("caminho") or "/")

            cur.execute("""
                SELECT 1
                  FROM galeria_juridico
                 WHERE caminho = %s
                   AND type    = %s
                   AND nome_arquivo = %s
                   AND NOT (
                        status = 'deletado'
                     OR (status = 'alterado' AND COALESCE(alteracao,'') = 'deletado')
                   )
                 LIMIT 1
            """, (caminho_final, tipo_dest, nome_base_final))
            existe = cur.fetchone() is not None
            if existe and fail_if_exists:
                return jsonify(status="erro",
                               codigo="nome_duplicado",
                               mensagem="Já existe um template com esse nome nesta pasta."), 409

            comp_json = json.dumps(
                [{"nome_arquivo": nome_base_final, "template_html": html_comp}],
                ensure_ascii=False
            )

            # Usa o usuário enviado ou o do template base
            usuario_final = usuario_req or base.get("usuario")

            nome_final = nome_unico(cnx, caminho_final, nome_base_final) if existe else nome_base_final

            new_id = executar("""
                INSERT INTO galeria_juridico
                       (nome_arquivo, status, alteracao,
                        id_origem,    usuario,
                        data_criacao, caminho, type,
                        template,     template_html, variaveis_valores, em_uso)
                VALUES (%s, %s, %s,
                        %s, %s,
                        NOW(), %s, %s,
                        %s, %s, %s, 0)
            """, (
                nome_final, "", None,
                id_base, usuario_final,
                caminho_final, tipo_dest,
                None, comp_json, None
            ))
        return jsonify(status="sucesso", mensagem="Documento criado com sucesso.", id_template=new_id, nome=nome_final, tipo=tipo_dest)
    except Exception:
        logger.exception("/salvar_como_novo: falha ao salvar como novo")
        return jsonify(status="erro", mensagem="Falha ao salvar como novo"), 500

@app.route("/inserir_componentes", methods=["POST"])
def inserir_componentes():
    """Insere, após um índice, componentes oriundos de outros templates (por id)."""
    data = request.get_json(force=True) or {}
    id_base = data.get("id_template_base")
    idx_comp_base = data.get("idx_comp_base")
    ids_inserir = data.get("ids_templates") or []

    if None in (id_base, idx_comp_base) or not ids_inserir:
        return jsonify(status="erro", mensagem="Parâmetros ausentes")
    try:
        idx_comp_base = int(idx_comp_base)
        ids_inserir = [int(i) for i in ids_inserir]
    except Exception:
        return jsonify(status="erro", mensagem="Índices/IDs inválidos")

    agora = marca_agora()
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            fmt = ",".join(["%s"] * len(ids_inserir))
            cur.execute(
                f"""SELECT id_template, nome_arquivo, template_html FROM galeria_juridico WHERE id_template IN ({fmt})""",
                tuple(ids_inserir))
            rows_ins = cur.fetchall()
            if len(rows_ins) != len(ids_inserir):
                return jsonify(status="erro", mensagem="Um ou mais templates para inserir não foram encontrados")

            novos_comp = []
            for r in rows_ins:
                try:
                    sub_components = json.loads(r["template_html"] or "[]")
                    if isinstance(sub_components, list):
                        novos_comp.extend(sub_components)
                    else:
                        novos_comp.append({"nome_arquivo": r["nome_arquivo"], "template_html": r["template_html"]})
                except Exception:
                    novos_comp.append({"nome_arquivo": r["nome_arquivo"], "template_html": r["template_html"] or ""})

            cur.execute("SELECT template_html FROM galeria_juridico WHERE id_template=%s", (id_base,))
            row_base = cur.fetchone()
            if not row_base:
                return jsonify(status="erro", mensagem="Template base não encontrado")

            comp_list = json.loads(row_base["template_html"] or "[]")
            if not (0 <= idx_comp_base < len(comp_list)):
                return jsonify(status="erro", mensagem="Índice do componente base fora de alcance")

            ponto_de_insercao = idx_comp_base + 1
            nova_lista = comp_list[:ponto_de_insercao] + novos_comp + comp_list[ponto_de_insercao:]

            cur.execute("""
                UPDATE galeria_juridico
                   SET template_html = %s,
                       data_alteracao = %s,
                       em_uso = FALSE
                 WHERE id_template = %s
            """, (json.dumps(nova_lista, ensure_ascii=False), agora, id_base))
            cnx.commit()

        return jsonify(status="sucesso", mensagem="Componentes inseridos com sucesso.", total_componentes=len(nova_lista), data_alteracao=agora)
    except Exception:
        logger.exception("/inserir_componentes: falha ao inserir")
        return jsonify(status="erro", mensagem="Falha ao inserir componentes"), 500

@app.route("/download_word", methods=["GET"])
def download_word():
    """Gera e baixa o DOCX final, aplicando variáveis salvas no template."""
    template_id = request.args.get("id", type=int)
    if not template_id:
        return jsonify(status="erro", mensagem="ID faltando"), 400
    try:
        htmls, vals, nome = _load_componentes_e_variaveis(template_id)
        if htmls is None:
            return jsonify(status="erro", mensagem="Template não encontrado"), 404
        html_fusao = _merge_componentes_com_variaveis(htmls, vals)
        docx_io = convert_html_to_docx(html_fusao)
        download_name = _safe_filename(nome, "docx")
        return send_file(
            docx_io,
            as_attachment=True,
            download_name=download_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception:
        logger.exception("/download_word: falha ao gerar DOCX")
        return jsonify(status="erro", mensagem="Falha ao gerar documento"), 500

@app.route("/unlock_template", methods=["POST"])
def unlock_template():
    """Libera lock (em_uso) de um template (editores/visualizadores)."""
    data = request.get_json(force=True) or {}
    id_template = data.get("id_template")
    if not id_template:
        return jsonify(status="erro", mensagem="O 'id_template' é obrigatório."), 400
    try:
        with get_connection() as cnx, cnx.cursor() as cur:
            cur.execute("""
                UPDATE galeria_juridico
                   SET em_uso = FALSE,
                       em_uso_modo = NULL,
                       em_uso_por  = NULL
                 WHERE id_template = %s
            """, (id_template,))
            if cur.rowcount == 0:
                cur.execute("SELECT id_template FROM galeria_juridico WHERE id_template = %s", (id_template,))
                if cur.fetchone() is None:
                    return jsonify(status="erro", mensagem=f"Template com ID {id_template} não encontrado."), 404
            cnx.commit()
        return jsonify(status="sucesso", mensagem=f"Template {id_template} foi desbloqueado com sucesso.")
    except Exception:
        logger.exception("/unlock_template: falha ao desbloquear")
        return jsonify(status="erro", mensagem="Erro interno ao desbloquear template."), 500

@app.route("/refresh_lock", methods=["POST"])
def refresh_lock():
    """Atualiza o timestamp do lock (mantém documento 'vivo' em uso)."""
    data = request.get_json(force=True) or {}
    id_template = data.get("id_template")
    if not id_template:
        return jsonify(status="erro", mensagem="ID do template ausente")

    agora = marca_agora()
    try:
        with get_connection() as cnx, cnx.cursor() as cur:
            cur.execute("""
                UPDATE galeria_juridico
                   SET data_alteracao = %s
                 WHERE id_template = %s AND em_uso = TRUE
            """, (agora, id_template))
            if cur.rowcount == 0:
                return jsonify(status="erro", mensagem="Documento não encontrado ou não está em uso")
            cnx.commit()
        return jsonify(status="sucesso", mensagem="Lock atualizado com sucesso.", data_alteracao=agora)
    except Exception:
        logger.exception("/refresh_lock: falha ao atualizar lock")
        return jsonify(status="erro", mensagem="Falha ao atualizar lock"), 500

@app.route("/download_pdf", methods=["GET"])
def download_pdf():
    """Gera e baixa o PDF final (conversão via LibreOffice/docx2pdf)."""
    template_id = request.args.get("id", type=int)
    if not template_id:
        return jsonify(status="erro", mensagem="ID faltando"), 400
    try:
        htmls, vals, nome = _load_componentes_e_variaveis(template_id)
        if htmls is None:
            return jsonify(status="erro", mensagem="Template não encontrado"), 404
        html_fusao = _merge_componentes_com_variaveis(htmls, vals)
        docx_io = convert_html_to_docx(html_fusao)
        pdf_io = _docx_bytes_to_pdf_bytes(docx_io)
        download_name = _safe_filename(nome, "pdf")
        return send_file(
            pdf_io,
            as_attachment=True,
            download_name=download_name,
            mimetype="application/pdf"
        )
    except Exception:
        logger.exception("/download_pdf: falha ao gerar PDF")
        return jsonify(status="erro", mensagem="Falha ao gerar PDF"), 500

@app.route("/linhas_pdf", methods=["POST"])
def linhas_pdf():
    """Gera 'âncoras' (primeira/última linha) por página do PDF total e por componente."""
    data = request.get_json(force=True) or {}

    def _to_int(v):
        try:
            return int(str(v).strip())
        except Exception:
            return None

    def _is_nonempty_str(s):
        return isinstance(s, str) and s.strip() != ""

    VAR_PATTERN = re.compile(r"\[([^\]\[]+)\]")

    def _vals_list_to_dict(lst):
        vals = {}
        if isinstance(lst, list):
            for it in lst:
                try:
                    n, v = it.get("nome"), it.get("valor")
                    if _is_nonempty_str(n) and v is not None:
                        vals[n] = str(v)
                except Exception:
                    pass
        return vals

    def _apply_vars(html_src, vals_dict):
        if not _is_nonempty_str(html_src):
            return ""
        def repl(m):
            k = m.group(1)
            val = vals_dict.get(k)
            return str(val) if _is_nonempty_str(val) else f"[{k}]"
        return VAR_PATTERN.sub(repl, html_src)

    def _wrap_html(body):
        return "<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>" + body + "</body></html>"

    def _pdf_pages_from_html(body_html):
        docx_io = convert_html_to_docx(body_html)
        pdf_io  = _docx_bytes_to_pdf_bytes(docx_io)
        pages   = _extract_pages_text_from_pdf(pdf_io, normalizar=True, manter_quebras=True)
        return pages

    def _first_last_lines(pages_text):
        first_lines, last_lines = [], []
        for ptxt in pages_text:
            raw_lines = [(ln or "").strip() for ln in (ptxt or "").splitlines()]
            lines = [ln for ln in raw_lines if ln]
            if not lines:
                first_lines.append("")
                last_lines.append("")
                continue
            def pick_first(ls):
                for s in ls:
                    if not s.isdigit():
                        return s
                return ls[0]
            def pick_last(ls):
                for s in reversed(ls):
                    if not s.isdigit():
                        return s
                return ls[-1]
            first_lines.append(pick_first(lines))
            last_lines.append(pick_last(lines))
        return first_lines, last_lines

    template_id = _to_int(data.get("id_template"))
    if not template_id:
        return jsonify(status="erro", mensagem="id_template faltando"), 400

    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("""
                SELECT nome_arquivo, template_html, variaveis_valores
                  FROM galeria_juridico
                 WHERE id_template = %s
            """, (template_id,))
            row = cur.fetchone()
            if not row:
                return jsonify(status="erro", mensagem="Template não encontrado"), 404

        try:
            comps_raw = json.loads(row["template_html"] or "[]")
            if not isinstance(comps_raw, list):
                comps_raw = [{"nome_arquivo": row["nome_arquivo"], "template_html": row["template_html"] or ""}]
        except Exception:
            comps_raw = [{"nome_arquivo": row["nome_arquivo"], "template_html": row["template_html"] or ""}]

        vals_db = []
        if row.get("variaveis_valores"):
            try:
                vals_db = json.loads(row["variaveis_valores"])
            except Exception:
                vals_db = []
        vals = _vals_list_to_dict(vals_db)

        if not comps_raw:
            return jsonify(status="erro", mensagem="Template sem componentes"), 400

        comps_aplicados = []
        for i, comp in enumerate(comps_raw):
            nome = comp.get("nome_arquivo") or f"Componente {i}"
            html_comp = comp.get("template_html") or ""
            html_ok = _apply_vars(html_comp, vals)
            comps_aplicados.append({"indice": i, "nome_arquivo": nome, "html": html_ok})

        full_body = "".join(c["html"] for c in comps_aplicados)
        pages_all = _pdf_pages_from_html(_wrap_html(full_body))
        if not pages_all:
            return jsonify(status="erro", mensagem="PDF não retornou páginas"), 500

        first_all, last_all = _first_last_lines(pages_all)

        componentes_out = []
        prev_until = 0
        for comp in comps_aplicados:
            i = comp["indice"]
            until_body = "".join(c["html"] for c in comps_aplicados[:i+1])
            pages_until = _pdf_pages_from_html(_wrap_html(until_body))
            start = prev_until
            end   = len(pages_until)
            prev_until = end
            first_slice = first_all[start:end]
            last_slice  = last_all[start:end]
            componentes_out.append({
                "indice": i,
                "nome_arquivo": comp["nome_arquivo"],
                "page_range": {"start": start, "end": end},
                "first_lines": first_slice,
                "last_lines": last_slice
            })

        return jsonify(
            status="sucesso",
            mensagem="Âncoras do PDF geradas com sucesso.",
            nome_arquivo=(row.get("nome_arquivo") if isinstance(row, dict) else None),
            total_paginas=len(pages_all),
            anchors_full={"first_lines": first_all, "last_lines": last_all},
            componentes=componentes_out
        )
    except Exception:
        logger.exception("/linhas_pdf: falha ao gerar âncoras")
        return jsonify(status="erro", mensagem="Falha ao gerar informações do PDF"), 500

@app.route("/transferir_variaveis", methods=["POST"])
def transferir_variaveis():
    """Transfere valores de variáveis do documento origem para destino, persistindo interseção."""
    try:
        data = request.get_json(force=True) or {}
    except Exception:
        return jsonify(status="erro", mensagem="JSON inválido"), 400

    id_origem  = data.get("id_origem")
    id_destino = data.get("id_destino")
    lista_vals = data.get("variaveis_valores")

    try:
        id_destino = int(id_destino)
    except Exception:
        return jsonify(status="erro", mensagem="id_destino inválido"), 400

    VAR_RE = re.compile(r'\[([^\]]+)\]')

    def extract_vars_from_template_html(template_html_raw):
        nomes = set()
        try:
            maybe_list = json.loads(template_html_raw or "[]")
            if isinstance(maybe_list, list):
                comps = maybe_list
            else:
                comps = [{"template_html": template_html_raw or ""}]
        except Exception:
            comps = [{"template_html": template_html_raw or ""}]
        for comp in comps:
            html_comp = comp.get("template_html") or ""
            soup = BeautifulSoup(html_comp, "html.parser")
            texto = soup.get_text(" ", strip=False)
            for v in VAR_RE.findall(texto or ""):
                if v:
                    nomes.add(v)
            for v in VAR_RE.findall(html_comp):
                if v:
                    nomes.add(v)
        return nomes

    def list_to_dict(vals_list):
        d = {}
        if isinstance(vals_list, list):
            for it in vals_list:
                try:
                    nome = it.get("nome")
                    val  = it.get("valor")
                    if isinstance(nome, str) and nome.strip() != "":
                        d[nome] = ("" if val is None else str(val))
                except Exception:
                    pass
        return d

    agora = marca_agora()
    try:
        with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
            cur.execute("""
                SELECT id_template, template_html, variaveis_valores
                  FROM galeria_juridico
                 WHERE id_template = %s
                   AND type IN ('template','individual')
            """, (id_destino,))
            dest = cur.fetchone()
            if not dest:
                return jsonify(status="erro", mensagem="Destino não encontrado"), 404

            vars_dest = extract_vars_from_template_html(dest.get("template_html"))

            dest_vals_dict = {}
            if dest.get("variaveis_valores"):
                try:
                    dest_vals_dict = list_to_dict(json.loads(dest["variaveis_valores"]))
                except Exception:
                    dest_vals_dict = {}

            if not lista_vals and id_origem:
                try:
                    id_origem = int(id_origem)
                except Exception:
                    id_origem = None
                if id_origem:
                    cur.execute("""
                        SELECT variaveis_valores
                          FROM galeria_juridico
                         WHERE id_template = %s
                           AND type IN ('template','individual')
                    """, (id_origem,))
                    orig = cur.fetchone()
                    if orig and orig.get("variaveis_valores"):
                        try:
                            lista_vals = json.loads(orig["variaveis_valores"])
                        except Exception:
                            lista_vals = None

            input_vals_dict = list_to_dict(lista_vals) if lista_vals else {}

            intersec_dict = {n: v for n, v in input_vals_dict.items() if n in vars_dest}

            merged = {**dest_vals_dict, **intersec_dict}
            merged = {n: merged[n] for n in merged.keys() if n in vars_dest}
            merged_list = [{"nome": n, "valor": merged[n]} for n in sorted(merged.keys())]

            cur.execute("""
                UPDATE galeria_juridico
                   SET variaveis_valores = %s,
                       data_alteracao    = %s
                 WHERE id_template = %s
            """, (json.dumps(merged_list, ensure_ascii=False), agora, id_destino))
            cnx.commit()

        return jsonify(
            status="sucesso",
            mensagem="Valores de variáveis transferidos com sucesso.",
            destino=id_destino,
            total_destino=len(vars_dest),
            atualizadas=len(intersec_dict),
            data_alteracao=agora
        )
    except Exception:
        logger.exception("/transferir_variaveis: falha ao transferir")
        return jsonify(status="erro", mensagem="Falha ao transferir variáveis"), 500

@app.route("/download_componentes_zip", methods=["POST"])
def download_componentes_zip():
	"""
	Gera um ZIP contendo cada componente exportado individualmente
	no formato solicitado (DOCX ou PDF).

	Payload (JSON):
	{
		"id_template": 123,
		"formato": "pdf" | "docx",
		"indices": [0,2,3]   # opcional; se ausente/ vazio => todos
	}
	"""
	try:
		data = request.get_json(force=True) or {}
	except Exception:
		return jsonify(status="erro", mensagem="JSON inválido"), 400

	template_id = data.get("id_template")
	formato = (data.get("formato") or "pdf").strip().lower()
	indices = data.get("indices") or []

	if not template_id:
		return jsonify(status="erro", mensagem="id_template faltando"), 400
	if formato not in ("pdf", "docx"):
		return jsonify(status="erro", mensagem="Formato inválido"), 400

	try:
		with get_connection() as cnx, cnx.cursor(dictionary=True) as cur:
			cur.execute("""
				SELECT nome_arquivo, template_html, variaveis_valores
				  FROM galeria_juridico
				 WHERE id_template = %s
			""", (template_id,))
			row = cur.fetchone()
			if not row:
				return jsonify(status="erro", mensagem="Template não encontrado"), 404

		try:
			comps_json = json.loads(row["template_html"] or "[]")
			if not isinstance(comps_json, list):
				comps_json = [{"nome_arquivo": row["nome_arquivo"], "template_html": row["template_html"] or ""}]
		except Exception:
			comps_json = [{"nome_arquivo": row["nome_arquivo"], "template_html": row["template_html"] or ""}]

		vals = {}
		if row.get("variaveis_valores"):
			try:
				for it in json.loads(row["variaveis_valores"]):
					n, v = it.get("nome"), it.get("valor")
					if n:
						vals[n] = v
			except Exception:
				pass

		total = len(comps_json)
		if not total:
			return jsonify(status="erro", mensagem="Template sem componentes"), 400

		# ✅ sem indices => todos
		if not indices:
			sel = list(range(total))
		else:
			try:
				sel = sorted({int(i) for i in indices})
			except Exception:
				return jsonify(status="erro", mensagem="Índices inválidos"), 400
			sel = [i for i in sel if 0 <= i < total]
			if not sel:
				return jsonify(status="erro", mensagem="Nenhum índice válido"), 400

		out_zip = io.BytesIO()
		with zipfile.ZipFile(out_zip, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
			for i in sel:
				comp = comps_json[i]
				nome_comp = comp.get("nome_arquivo") or f"Componente_{i+1}"
				html_raw  = comp.get("template_html") or ""
				html_ok   = _apply_variables_to_html(html_raw, vals)

				docx_io = convert_html_to_docx(html_ok)
				if formato == "pdf":
					file_io = _docx_bytes_to_pdf_bytes(docx_io)
					ext = "pdf"
				else:
					file_io = docx_io
					ext = "docx"

				filename = _safe_filename(nome_comp, ext)
				zf.writestr(filename, file_io.getvalue())

		out_zip.seek(0)
		zip_name = _safe_filename(f"{row.get('nome_arquivo')}_componentes", "zip")
		return send_file(
			out_zip,
			as_attachment=True,
			download_name=zip_name,
			mimetype="application/zip"
		)
	except Exception:
		logger.exception("/download_componentes_zip: falha ao gerar ZIP")
		return jsonify(status="erro", mensagem="Falha ao gerar ZIP"), 500

# === TASKS / SCHEDULE ===
def unlock_stale_templates(limite_minutos=10):
    """Desbloqueia templates com lock antigo (inativos por 'limite_minutos')."""
    try:
        with get_connection() as cnx, cnx.cursor() as cur:
            cur.execute("""
                UPDATE galeria_juridico
                   SET em_uso = FALSE,
                       em_uso_modo = NULL,
                       em_uso_por  = NULL
                 WHERE em_uso = TRUE
                   AND data_alteracao IS NOT NULL
                   AND data_alteracao < NOW() - INTERVAL %s MINUTE
            """, (limite_minutos,))
            liberados = cur.rowcount
            cnx.commit()
        if liberados:
            logger.info("unlock_stale_templates: desbloqueados automaticamente: %s", liberados)
    except Exception:
        logger.exception("unlock_stale_templates: falha no desbloqueio automático")

# === MAIN ===
if __name__ == "__main__":
    def run_schedule():
        """Loop do scheduler para executar tarefas pendentes."""
        while True:
            try:
                schedule.run_pending()
            except Exception:
                logger.exception("run_schedule: exceção ao executar tarefas agendadas")
            time.sleep(1)

    schedule.every(10).seconds.do(unlock_stale_templates)

    t = threading.Thread(target=run_schedule, name="scheduler", daemon=True)
    t.start()

    logger.info("Servidor iniciando em 0.0.0.0:5001")
    app.run(host="0.0.0.0", port=5001)